from tkinter import *
from tkinter import ttk, messagebox
from docxtpl import DocxTemplate
import datetime
import time
import os
import mysql.connector as sql
from docx import Document

tk = Tk()
tk.geometry("1250x720")
tk.config(bg="#F5F5F5")
tk.title("Delivery Slip - Gayathri Tex")
tk.configure(bg='white')

mycon = sql.connect(host="localhost", user="root", passwd="Gayathri@2404", database="gtex_db")
cursor = mycon.cursor()

show = ("SELECT COUNT(*) FROM delivery_slip")
cursor.execute(show)
slip_no = cursor.fetchall()
slip_no = int(slip_no[0][0]) + 1

def clear_item():
    particulars_entry.delete(0, END)
    quantity_entry.delete(0, END)
    quantity_entry.insert(0, "0")

delivery_list = []

def add_item():
    particulars = particulars_entry.get()
    quantity = int(quantity_entry.get())
    delivery_item = [particulars, quantity]
    tree.insert('', 0, values=delivery_item)
    clear_item()
    delivery_list.append(delivery_item)

def generate_delivery_slip():
    name = name_entry.get()
    address = address_entry.get()
    phone = phone_number_entry.get()
    gst = gst_entry.get()
    current_date = datetime.date.today().strftime("%Y-%m-%d")
    current_time = time.strftime("%H:%M:%S")

    insert_slip = ("INSERT INTO delivery_slip (name, address, phone_number, gst_number, delivery_date) VALUES (%s, %s, %s, %s, %s)")
    slip_data = (name, address, phone, gst, current_date)
    cursor.execute(insert_slip, slip_data)
    mycon.commit()

    slip_id = cursor.lastrowid

    for item in delivery_list:
        insert_details = ("INSERT INTO delivery_slip_details (delivery_slip_no, particulars, quantity) VALUES (%s, %s, %s)")
        details_data = (slip_id, item[0], item[1])
        cursor.execute(insert_details, details_data)
        mycon.commit()

    generate_pdf(slip_id)

    messagebox.showinfo("Delivery Slip Generated", "Delivery Slip Generated")

def generate_pdf(slip_id):
    doc = Document()
    doc.add_heading('Delivery Slip', level=1)

    # Fetch slip data
    cursor.execute("SELECT * FROM delivery_slip WHERE delivery_slip_no=%s", (slip_id,))
    slip_data = cursor.fetchone()

    doc.add_paragraph(f"Name: {slip_data[1]}")
    doc.add_paragraph(f"Address: {slip_data[2]}")
    doc.add_paragraph(f"Phone Number: {slip_data[3]}")
    doc.add_paragraph(f"GST Number: {slip_data[4]}")
    doc.add_paragraph(f"Delivery Date: {slip_data[5]}")

    doc.add_heading('Items:', level=2)

    # Fetch slip details
    cursor.execute("SELECT * FROM delivery_slip_details WHERE delivery_slip_no=%s", (slip_id,))
    slip_details = cursor.fetchall()

    for detail in slip_details:
        doc.add_paragraph(f"Particulars: {detail[2]}, Quantity: {detail[3]}")

    # Save the document
    file_path = os.path.join("E:\\gtex_billing\\Billing_Software\\Delivery Slips", f"delivery_slip_{slip_id}.docx")
    doc.save(file_path)

    # Convert to PDF
    os.system(f"libreoffice --headless --convert-to pdf {file_path}")

def clear_delivery_slip():
    name_entry.delete(0, END)
    address_entry.delete(0, END)
    phone_number_entry.delete(0, END)
    gst_entry.delete(0, END)

    particulars_entry.delete(0, END)
    quantity_entry.delete(0, END)
    quantity_entry.insert(0, "0")

    for item in tree.get_children():
        tree.delete(item)

a = Label(tk, text="GAYATHRI TEX - ERODE", font=("Arial", 20, "bold"), bg="white", fg="#1A374D").place(x=400, y=20)
slip_number_label = Label(tk, text="Delivery Slip Number : DS-2425-{number:06}".format(number=slip_no), font=("Arial", 12, "bold"), bg="white", fg="#1A374D").place(x=900, y=20)

name_label = Label(tk, text="Name", font=("Arial", 12, "bold"), bg="white", fg="#1A374D").place(x=120, y=100)
name_entry = Entry(tk, font=("Arial", 12), bd=3)
name_entry.place(x=50, y=130)

address_label = Label(tk, text="Address", font=("Arial", 12, "bold"), bg="white", fg="#1A374D").place(x=440, y=100)
address_entry = Entry(tk, font=("Arial", 12), bd=3)
address_entry.place(x=375, y=130)

phone_number_label = Label(tk, text="Phone number", font=("Arial", 12, "bold"), bg="white", fg="#1A374D").place(x=740, y=100)
phone_number_entry = Entry(tk, font=("Arial", 12), bd=3)
phone_number_entry.place(x=700, y=130)

gst_label = Label(tk, text="GST number", font=("Arial", 12, "bold"), bg="white", fg="#1A374D").place(x=1068, y=100)
gst_entry = Entry(tk, font=("Arial", 12), bd=3)
gst_entry.place(x=1025, y=130)

particulars_label = Label(tk, text="Particulars", font=("Arial", 12, "bold"), bg="white", fg="#1A374D").place(x=440, y=180)
particulars_entry = Entry(tk, font=("Arial", 12), bd=3)
particulars_entry.place(x=375, y=210)

quantity_label = Label(tk, text="Quantity", font=("Arial", 12, "bold"), bg="white", fg="#1A374D").place(x=740, y=180)
quantity_entry = Spinbox(tk, from_=1, to=100, font=("Arial", 12), bd=3)
quantity_entry.place(x=700, y=210)

add_item_button = Button(tk, text="Add item", font=("Arial", 10, "bold"), bg="#1A374D", fg="#F5F5F5", command=add_item)
add_item_button.place(x=1110, y=210)

columns = ('particulars', 'quantity')
tree = ttk.Treeview(tk, columns=columns, show="headings")
tree.heading('particulars', text='Particulars')
tree.heading('quantity', text='Quantity')
tree.place(x=400, y=250)

save_slip_button = Button(tk, text="Generate Delivery Slip", font=("Arial", 10, "bold"), bg="#1A374D", fg="#F5F5F5", width=30, height=2, command=generate_delivery_slip)
save_slip_button.place(x=480, y=500)
new_slip_button = Button(tk, text="Clear Delivery Slip", font=("Arial", 10, "bold"), bg="#1A374D", fg="#F5F5F5", width=30, height=2, command=clear_delivery_slip)
new_slip_button.place(x=480, y=540)
edit_slip_button = Button(tk, text="Edit Delivery Slip", font=("Arial", 10, "bold"), bg="#1A374D", fg="#F5F5F5", width=30, height=2)
edit_slip_button.place(x=480, y=580)

tk.mainloop()
